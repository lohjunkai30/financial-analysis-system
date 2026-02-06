import pandas as pd
import numpy as np
from sqlalchemy import create_engine
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

# 1. Database Connection
load_dotenv()
DB_URL = os.getenv("DB_URL")
engine = create_engine(DB_URL)

def run_analysis(symbol, document, target_dir):
    # 2. Extract Data
    query = f"SELECT * FROM stock_data WHERE symbol = '{symbol}' ORDER BY action_date ASC"
    df = pd.read_sql(query, engine)
    if df.empty:
        return "No data"

    # 3. Generating the Indicators     
    # 3a. Overnight Gaps
    df["prev_close"] = df["close_price"].shift(1)
    df["overnight_gap"] = df["open_price"] - df["prev_close"]

    # 3b. Deriving Volatility
    volatility_formula = 100*(df["high_price"] - df["low_price"])/df["close_price"]
    df["volatility"] = np.where(volatility_formula>3, "High", "Low")

    # 3c. Setting up the Golden and Death Cross
    df["SMA50"] = df["close_price"].rolling(window=50).mean()
    df["SMA200"] = df["close_price"].rolling(window=200).mean()

    #FIX 1 (Remove the first 199 rows where SMA200 is NaN)
    df.dropna(inplace=True)

    # Identifying the "Crossovers" (1.0 = Golden Cross (Buy), -1.0 = Death Cross (Sell)).
    df["signal"] = np.where(df["SMA50"] > df["SMA200"], 1.0, 0.0)
    df["crossover"] = df["signal"].diff()
    crosses = df[(df["crossover"] !=0) & (df.index != df.index[0])].copy()

    # 3d. Custom Decision Engine
    last_row = df.iloc[-1]
    overnight_gap = last_row["overnight_gap"]
    volatility = last_row["volatility"]
    is_bullish = last_row["SMA50"] > last_row["SMA200"]

    if crosses.empty:
        if is_bullish:
            trend_type = "Persistent Bullish"
            verdict = "Hold/Monitor" 
        else:
            trend_type = "Persistently Bearish"
            verdict = "Exit/Neutral"
    else:
        last_cross_row = crosses.iloc[-1]
        if last_cross_row["crossover"] == 1:
            trend_type = "Golden Cross Detected"
            if overnight_gap >=0 and volatility == "High":
                verdict = "Speculative Buy"
            elif overnight_gap >=0 and volatility == "Low":
                verdict = "Strong Buy"
            elif overnight_gap <0:
                verdict = "Hold/Monitor"
        else:
            trend_type = "Death Cross Detected"
            if overnight_gap >=0 and volatility == "High":
                verdict = "Speculative Sell"
            elif overnight_gap >=0 and volatility == "Low":
                verdict = "Strong Sell"
            elif overnight_gap <0:
                verdict = "Exit/Neutral"

    # 3e. Adding into the document

    document.add_heading(f"Analysis for {symbol}", level=1)
    para1 = document.add_paragraph()
    para1.add_run(f"\nTrend Type: {trend_type}")
    para1.add_run(f"\nOvernight Gap: ${overnight_gap:.2f}")
    para1.add_run(f"\nVolatility: {volatility}")
    para1.add_run(f"\nFinal Verdict: {verdict}")

    # 7. Visualization
    print(f"Generating charts for {symbol}...")
    plt.figure(figsize=(12,6))
    
    # 7a. Panel 1: Price and Moving Averages

    #Slicing my data to reflect only the most recent 120 days to be shown on the chart
    df_sliced = df.tail(250)


    plt.subplot(2, 1, 1)
    plt.plot(df_sliced["action_date"], df_sliced["close_price"], label="Price", alpha=0.4)
    plt.plot(df_sliced["action_date"], df_sliced["SMA50"], label="50-Day MA", color="orange")
    plt.plot(df_sliced["action_date"], df_sliced["SMA200"], label="200-Day MA", color="blue")
    plt.title(f"{symbol} Trend Analysis")
    plt.legend()
    plt.grid(True, alpha=0.3)

    # 7b. Panel 2: The "Gaps" (or Volume)
    plt.subplot(2, 1, 2)
    conditions = [
        (df_sliced["overnight_gap"] > 0),
        (df_sliced["overnight_gap"] <0)
    ]
    choices = [
        "green", "red"
    ]
    colors = np.select(conditions, choices, default = "grey")
    plt.bar(df_sliced["action_date"], df_sliced["overnight_gap"], color = colors, alpha = 0.3)
    plt.axhline(y=0, color= "black", linewidth = 0.5)
    plt.title("Overnight Gaps ($)")
    plt.grid(True, alpha = 0.3)
    plt.xticks(rotation =45)

    # 7c. Final touches to the Report Picture
    plt.tight_layout()
    file_name = (f"{symbol}_report.png")
    save_path = os.path.join(target_dir, file_name)
    plt.savefig(save_path)
    document.add_picture(save_path, width = Inches(5.5))
    plt.close()
    print(f"Chart saved to {save_path} and added to report")

    


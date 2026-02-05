import os
import sys

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from scraping import fetch_and_save_stock
from analysis import run_analysis
from datetime import date, timedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# 1. Defining the Portfolio List
portfolio = ["AAPL", "TSLA", "NVDA", "AMD", "AMZN", "MSFT", "GOOGL"]


def main():
    print("Starting Process")

    # Phase 1: Creating the Folder and Report for that day
    # Phase 1a: Creating the Folder directory
    today = date.today()
    dailyreport_folder = f"{today} report"
    target_dir = os.path.join("reports", dailyreport_folder) #target_dir is now a folder path reports/today report
    os.makedirs(target_dir, exist_ok = True) #creates the above folder path

    # Phase 1b: Creating the Report document
    document = Document()
    doc_name = f"{today}_report.docx"
    doc_folder = target_dir
    doc_path = os.path.join(doc_folder, doc_name)
    heading0 = document.add_heading("Daily Financial Strategy Report", 0)
    heading0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_date = document.add_paragraph(f"Generated On: {today}", )
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    # Phase 1: Updating my SQL Database
    print("\n---Updating Database---")
    # Generating Live Dates for dynamic scraping
    today_latest = today + timedelta(days=1)
    today_latest_str = today_latest.strftime('%Y-%m-%d')
    start_date = today - timedelta(days=500)
    start_date_str = start_date.strftime('%Y-%m-%d')
    for i in portfolio:
        try:
            fetch_and_save_stock(i, start_date_str, today_latest_str)
        except Exception as e:
            print(f"Scraping for {i} failed due to {e}")
    
    # Phase 2: Generating the Analysis
    print("\n---Generating PNG Reports---")
    for i in portfolio:
        try:
            run_analysis(i, document, target_dir)
        except Exception as e:
            print(f"Unable to generate report for {i} due to {e}")

    # Phase 3: Saving the finalised generated report
    document.save(doc_path)
        
    # Phase 4: Indicating End of the Function
    print("\n All Analysis Completed")

if __name__ == "__main__":
    main()
